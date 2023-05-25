import openai
import config
import pandas as pd
import yfinance as yf
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import ttk

# set openai api key
openai.api_key = config.API_KEY

# Set default start and end dates
default_start_date = "2020-01-01"
default_end_date = datetime.today().strftime('%Y-%m-%d')

# Get user input


def main():
    # read report
    def read_report(ticker_symbol):
        filename = ticker_symbol + '_report.docx'
        document = Document(filename)
        report = ""
        for para in document.paragraphs:
            report += para.text

        return report

    # analyze report
    def analyze_report(ticker_symbol):

        report = read_report(ticker_symbol)

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": "You are a helpful investment advisor."},
                      {"role": "user", "content": f"This is a report on a company: {report} Do you think this company is a good investment?"}],
            temperature=0.2,
            max_tokens=300,
            top_p=1,
            frequency_penalty=0.0,
            presence_penalty=0.0)
        return response['choices'][0]['message']['content']

    def output_report(ticker_symbol):

        response = analyze_report(ticker_symbol)

        # Create a new document
        doc = Document()
        doc.add_heading(f'Chatgpt Report for {ticker_symbol}', 0)
        doc.add_paragraph(response)
        doc.save(f"{ticker_symbol}_Chatgpt_report.docx")

    #  Create root window
    root = tk.Tk()

    # Create frames
    input_frame = tk.Frame(root)
    input_frame.pack(side="top", fill="x")

    output_frame = tk.Frame(root)
    output_frame.pack(side="bottom", fill="both", expand=True)

    # User Input
    ticker_label = tk.Label(input_frame, text="Enter Ticker Symbol:")
    ticker_entry = tk.Entry(input_frame)
    # start_date_label = tk.Label(input_frame, text="Start date:")
    # start_date_entry = tk.Entry(input_frame)
    # end_date_label = tk.Label(input_frame, text="End date:")
    # end_date_entry = tk.Entry(input_frame)

    # Set default values for start and end date entries
    # start_date_entry.insert(tk.END, default_start_date)
    # end_date_entry.insert(tk.END, default_end_date)

    # Position widgets using grid
    ticker_label.grid(row=0, column=0)
    ticker_entry.grid(row=0, column=1)
    # start_date_label.grid(row=1, column=0)
    # start_date_entry.grid(row=1, column=1)
    # end_date_label.grid(row=2, column=0)
    # end_date_entry.grid(row=2, column=1)

    fetch_button = ttk.Button(
        input_frame, text="Get Report", command=lambda: output_report(ticker_entry.get()))
    fetch_button.grid(row=3, column=0, columnspan=2)

    root.mainloop()


if __name__ == "__main__":
    main()
